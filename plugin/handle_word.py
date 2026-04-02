import os
import re
import tempfile
import uuid
import copy
from docx import Document
from docx.text.paragraph import Paragraph
from docx.shared import Pt
from langdetect import detect
import argostranslate.translate
import argostranslate.sbd
import win32com.client
import pythoncom 
from concurrent.futures import ThreadPoolExecutor, as_completed # THÊM ĐA LUỒNG

# --- 1.  BYPASS STANZA 
def offline_split_sentences(self, text):
    temp_text = re.sub(r'(?<=\d)\.(?=\d)', '___DOT___', text)
    sentences = re.split(r'(?<=[。！？])|(?<=[.!?])\s+', temp_text)
    return [s.replace('___DOT___', '.').strip() for s in sentences if s and s.strip()]

for name, obj in vars(argostranslate.sbd).items():
    if isinstance(obj, type) and hasattr(obj, 'split_sentences'):
        setattr(obj, 'split_sentences', offline_split_sentences)

def fix_automotive_terms(text):
    dict_replacements = {
        "Mental Driving": "Intelligent Driving",
        "mental driving": "intelligent driving",
        "Line Control": "Drive-by-wire",
        "line control": "drive-by-wire",
        "Seo Ning Ning": "Xu Ningning", 
        "Medium-compatible": "Basic", 
    }
    for wrong, right in dict_replacements.items():
        text = text.replace(wrong, right)
    return text

def get_translator(from_code, to_code):
    installed_languages = argostranslate.translate.get_installed_languages()
    lang_dict = {lang.code: lang for lang in installed_languages}
    if from_code in lang_dict and to_code in lang_dict:
        return lang_dict[from_code].get_translation(lang_dict[to_code])
    return None

def smart_detect_lang(text):
    if not text or not text.strip(): return 'unknown'
    if re.search(r'[\u4e00-\u9fff]', text): return 'zh'
    vi_chars = r'[áàảãạăắằẳẵặâấầẩẫậéèẻẽẹêếềểễệíìỉĩịóòỏõọôốồổỗộơớờởỡợúùủũụưứừửữựýỳỷỹỵđ]'
    if re.search(vi_chars, text.lower()): return 'vi'
    try:
        lang = detect(text)
        return 'zh' if lang.startswith('zh') else lang
    except: return 'unknown'

def do_translation(text, source_code, target_lang_code):
    if source_code == target_lang_code or source_code == 'unknown': return text
    translator = get_translator(source_code, target_lang_code)
    
    trans_text = text
    if translator: 
        trans_text = translator.translate(text)
    elif source_code != 'en' and target_lang_code != 'en':
        trans_to_en = get_translator(source_code, 'en')
        trans_to_target = get_translator('en', target_lang_code)
        if trans_to_en and trans_to_target:
            trans_text = trans_to_target.translate(trans_to_en.translate(text))
            
    return fix_automotive_terms(trans_text)

def translate_text_block(text, target_lang_code):
    hex_pattern = re.compile(r'\b[A-Fa-f0-9]{10,}\b')
    protected_items = []
    
    def mask_repl(match):
        protected_items.append(match.group(0))
        return f" MASK{len(protected_items)-1} "
        
    text_hidden = hex_pattern.sub(mask_repl, text)
    gap_pattern = re.compile(r'((?:[.\-_][ \t]?){4,}|[ \t]{4,})')
    
    lines = text_hidden.splitlines(keepends=False)
    result_lines = []
    
    for line in lines:
        if not line.strip():
            result_lines.append("")
            continue
            
        chunks = gap_pattern.split(line)
        trans_chunks = []
        
        for chunk in chunks:
            if not chunk or gap_pattern.fullmatch(chunk):
                trans_chunks.append(chunk)
                continue
                
            clean_chunk = chunk.strip()
            
            prefix = ""
            text_to_trans = clean_chunk
            num_prefix_match = re.match(r'^([\d\.\s]+)(.*)', clean_chunk)
            if num_prefix_match:
                prefix = num_prefix_match.group(1)
                text_to_trans = num_prefix_match.group(2)

            text_without_mask = re.sub(r'MASK\d+', '', text_to_trans).strip()
            
            if not text_without_mask or not re.search(r'[a-zA-Z\u4e00-\u9fffÀ-ỹ]', text_without_mask):
                trans_chunks.append(chunk)
                continue
            
            if len(text_to_trans) > 1500:
                text_to_trans = text_to_trans[:1500] + "..."
                
            try:
                source_code = smart_detect_lang(text_to_trans)
                if source_code == 'unknown': source_code = 'en'
                
                trans_text = do_translation(text_to_trans, source_code, target_lang_code)
                
                left_space = chunk[:len(chunk) - len(chunk.lstrip())]
                right_space = chunk[len(chunk.rstrip()):]
                
                trans_chunks.append(f"{left_space}{prefix}{trans_text}{right_space}")
                
            except Exception as e:
                print(f"[DEBUG] Lỗi dịch chunk: {e}")
                trans_chunks.append(chunk)
                
        result_lines.append("".join(trans_chunks))
        
    final_text = '\n'.join(result_lines)
    
    for i, item in enumerate(protected_items):
        final_text = final_text.replace(f" MASK{i} ", item).replace(f"MASK{i}", item)
        
    return final_text

def get_all_paragraphs(doc):
    paras = []
    for p in doc.paragraphs: paras.append(p)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs: paras.append(p)
    for section in doc.sections:
        for header_type in [section.header, section.first_page_header, section.even_page_header]:
            if header_type and not header_type.is_linked_to_previous:
                for p in header_type.paragraphs: paras.append(p)
                for table in header_type.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for p in cell.paragraphs: paras.append(p)
        for footer_type in [section.footer, section.first_page_footer, section.even_page_footer]:
            if footer_type and not footer_type.is_linked_to_previous:
                for p in footer_type.paragraphs: paras.append(p)
                for table in footer_type.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for p in cell.paragraphs: paras.append(p)
    for txbx in doc.element.xpath('.//*[local-name()="txbxContent"]'):
        for p_xml in txbx.xpath('.//*[local-name()="p"]'): paras.append(Paragraph(p_xml, doc))
    for txbx in doc.element.xpath('.//*[local-name()="textbox"]'):
        for p_xml in txbx.xpath('.//*[local-name()="p"]'): paras.append(Paragraph(p_xml, doc))
    return paras

def convert_doc_to_docx(input_path):
    print(f"[DEBUG] Converting .doc file to .docx using MS Word...")
    pythoncom.CoInitialize() 
    word = None
    doc = None
    try:
        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False 
        word.DisplayAlerts = False 
        
        abs_in = os.path.abspath(input_path)
        temp_dir = tempfile.gettempdir()
        unique_filename = f"bosch_trans_temp_{uuid.uuid4().hex}.docx"
        abs_out = os.path.join(temp_dir, unique_filename)
        
        doc = word.Documents.Open(abs_in, ConfirmConversions=False, ReadOnly=True)
        doc.SaveAs2(abs_out, FileFormat=16) 
        return abs_out
    except Exception as e:
        print(f"[DEBUG] Error during COM conversion: {e}")
        raise e
    finally:
        if doc:
            try: doc.Close(SaveChanges=False)
            except: pass
        if word:
            try: word.Quit()
            except: pass

def replace_para_text_vip_pro(para, translated_text, original_text):
    if not para.text.strip(): return
    
    text_runs = [run for run in para.runs if run.text.strip()]
    if not text_runs: return
    
    host_run = max(text_runs, key=lambda r: len(r.text.strip()))
    
    for run in text_runs:
        run.text = ""
        
    host_run.text = translated_text
    
    ratio = len(translated_text) / max(len(original_text), 1)
    if ratio > 1.15:
        for run in para.runs:
            if run.font.size:
                current_pt = run.font.size.pt
                new_pt = max(8.0, current_pt - min(3.0, (ratio - 1) * 2))
                run.font.size = Pt(new_pt)

    try:
        para.paragraph_format.keep_with_next = False
        para.paragraph_format.page_break_before = False
        para.paragraph_format.keep_together = False
        if para.paragraph_format.line_spacing is None or para.paragraph_format.line_spacing > 1.5:
            para.paragraph_format.line_spacing = 1.15
    except:
        pass
    
def remove_personal_info_warning(doc):
    try:
        settings = doc.settings.element
        warning_tags = settings.xpath('.//*[local-name()="removePersonalInformation"]')
        for tag in warning_tags:
            tag.getparent().remove(tag)
    except: pass

# Hàm worker đa luồng
def worker_translate_para(original_text, target_lang_code):
    try:
        return translate_text_block(original_text, target_lang_code)
    except Exception:
        return original_text

def process_core(input_docx, output_file, target_lang_code, progress_callback, cancel_event=None, log_callback=None):
    def log(msg):
        if log_callback: log_callback(msg)
        else: print(msg)

    doc = Document(input_docx)
    all_paras = get_all_paragraphs(doc)
    
    items_to_translate = [p for p in all_paras if p.text.strip()]
    total_items = len(items_to_translate)
    
    if total_items == 0: return "No content found to translate."

    log(f"Total paragraphs to translate: {total_items}. Kích hoạt Parallel Processing...")

    current_item = 0
    
    # -- CODE MỚI: TỰ ĐỘNG TÍNH TOÁN SỐ LUỒNG CHO WORD --
    import os
    cpu_cores = os.cpu_count() or 4 
    
    # Công thức: Bằng 75% số nhân CPU (Tối thiểu 2 luồng, Tối đa 12 luồng)
    # Tránh bung quá nhiều luồng làm nghẽn CPU khi parse cấu trúc XML của Word
    MAX_WORKERS = max(2, min(12, int(cpu_cores * 0.75))) 
    
    print(f"[DEBUG] System has {cpu_cores} CPU cores. Auto-set MAX_WORKERS = {MAX_WORKERS} for Word doc")
    # ---------------------------------------------------
    
    with ThreadPoolExecutor(max_workers=MAX_WORKERS) as executor:
        future_to_para = {}
        for para in items_to_translate:
            future = executor.submit(worker_translate_para, para.text, target_lang_code)
            future_to_para[future] = para

        for future in as_completed(future_to_para):
            if cancel_event and cancel_event.is_set():
                log("🛑 User requested cancellation. Saving temporary file...")
                executor.shutdown(wait=False, cancel_futures=True)
                remove_personal_info_warning(doc)
                doc.save(output_file)
                return f"Translation cancelled!\nPartial file saved at: {os.path.basename(output_file)}"

            para = future_to_para[future]
            try:
                translated_text = future.result()
                if translated_text != para.text:
                    # Việc gán text ngược lại vào Object của Docx bắt buộc tuần tự để an toàn
                    replace_para_text_vip_pro(para, translated_text, para.text)
            except Exception as e:
                print(f"[ERROR] Dịch đoạn văn thất bại: {e}")
                
            current_item += 1
            progress_callback(current_item / total_items)

            if current_item % 50 == 0 or current_item == total_items:
                percent = (current_item / total_items) * 100
                log(f"Tiến độ: {current_item}/{total_items} ({percent:.1f}%) - Đang chạy max tốc độ...")

    remove_personal_info_warning(doc)
    doc.save(output_file)
    return f"Translation completed!\nFile saved at: {os.path.basename(output_file)}"

def process(input_path, output_dir, target_lang_code, progress_callback, cancel_event=None, log_callback=None):
    original_input = input_path
    temp_docx_path = None
    
    if input_path.lower().endswith('.doc'):
        temp_docx_path = convert_doc_to_docx(input_path)
        input_path = temp_docx_path

    base_name = os.path.basename(original_input)
    file_name, _ = os.path.splitext(base_name)
    output_filename = f"{file_name}_Translated_{target_lang_code.upper()}.docx"
    output_file = os.path.join(output_dir, output_filename)

    result = process_core(input_path, output_file, target_lang_code, progress_callback, cancel_event, log_callback)

    if temp_docx_path and os.path.exists(temp_docx_path):
        os.remove(temp_docx_path)
        
    return result